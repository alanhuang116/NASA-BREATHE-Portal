"""Generate a detailed budget table for BREATHE Section 7.4 Citizen Science."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import re

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.90)
    section.right_margin = Cm(1.90)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15


def add_breathe_header():
    header_para = doc.add_paragraph()
    parts = [
        ('BREATHE: ', True, False),
        ('BR', True, True), ('idging ', True, False),
        ('E', True, True), ('ducation on ', True, False),
        ('A', True, True), ('tmosphere, ', True, False),
        ('T', True, True), ('echnology, and ', True, False),
        ('HE', True, True), ('alth', True, False),
    ]
    for text, bold, underline in parts:
        r = header_para.add_run(text)
        r.bold = bold
        r.underline = underline
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'
    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_after = Pt(6)
    pPr = border_para._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')


def add_heading(text, size=Pt(12)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = size
    r.font.name = 'Times New Roman'
    return p


def add_body(text, space_after=Pt(4)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.15
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'
    return p


def shade_row(row, color):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color, qn('w:val'): 'clear'
        })
        shading.append(shading_elm)


def set_cell(cell, text, bold=False, size=Pt(9), align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text))
    r.font.size = size
    r.font.name = 'Times New Roman'
    r.bold = bold


def add_budget_table(title, columns, rows, totals=None, note=None):
    """Add a formatted budget table."""
    add_heading(title, size=Pt(11))

    ncols = len(columns)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, col in enumerate(columns):
        set_cell(hdr.cells[i], col, bold=True, size=Pt(9),
                 align='center' if i > 0 else 'left')
    shade_row(hdr, '003366')
    for cell in hdr.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        is_subtotal = row_data[0].startswith('Subtotal') or row_data[0].startswith('TOTAL')
        for i, val in enumerate(row_data):
            align = 'right' if i >= 2 else 'left'
            set_cell(row.cells[i], val, bold=is_subtotal, size=Pt(9), align=align)
        if is_subtotal:
            shade_row(row, 'E8EEF4')

    # Totals row
    if totals:
        row = table.add_row()
        for i, val in enumerate(totals):
            align = 'right' if i >= 2 else 'left'
            set_cell(row.cells[i], val, bold=True, size=Pt(9), align=align)
        shade_row(row, '003366')
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(255, 255, 255)

    if note:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(note)
        r.font.size = Pt(8)
        r.font.name = 'Times New Roman'
        r.italic = True
        r.font.color.rgb = RGBColor(100, 100, 100)


# ============================================================
# DOCUMENT
# ============================================================
add_breathe_header()
add_heading('Section 7.4 — Citizen Science and Community Monitoring: Detailed Budget', size=Pt(12))

add_body(
    'This budget covers all hardware, deployment, maintenance, and operational costs for '
    'BREATHE\u2019s tiered citizen science monitoring activities across the three project stages. '
    'Sensor pricing is based on current (2025\u201326) manufacturer list prices. Personnel costs '
    'for citizen science coordination are included in the overall project personnel budget and '
    'are not duplicated here.'
)

# ============================================================
# STAGE 1
# ============================================================
add_budget_table(
    'Stage 1: Sensor-Based Monitoring in Pilot Communities (Years 1\u20132)',
    ['Item', 'Description', 'Unit Cost', 'Qty', 'Total'],
    [
        # --- Fixed outdoor sensors ---
        ['', 'A. FIXED OUTDOOR SENSOR NETWORK', '', '', ''],
        ['PurpleAir Flex (outdoor PM2.5)',
         'WiFi-enabled, dual-laser PM2.5 sensor for continuous outdoor monitoring; '
         'data auto-uploaded to PurpleAir map and BREATHE platform',
         '$289', '30', '$8,670'],
        ['Outdoor power supply',
         'Weatherproof power adapter for PurpleAir Flex (not included with sensor)',
         '$45', '30', '$1,350'],
        ['Mounting hardware & enclosure',
         'Pole/wall mount brackets, zip ties, weatherproof junction box, UV-resistant cable',
         '$35', '30', '$1,050'],
        ['MicroSD cards (32 GB)',
         'Local data backup on each sensor (not included with PurpleAir Flex)',
         '$10', '30', '$300'],
        ['Ethernet cable & adapters',
         'For sites where WiFi is unreliable; Cat6 outdoor-rated cable (50 ft) + adapter',
         '$25', '10', '$250'],
        ['Subtotal A: Fixed Outdoor Sensors', '', '', '', '$11,620'],

        # --- Portable sensors ---
        ['', 'B. PORTABLE SENSORS FOR FIELD INVESTIGATIONS', '', '', ''],
        ['AirBeam Mini (portable PM2.5)',
         'Handheld, battery-powered (14 hr), Bluetooth-connected PM2.5 monitor '
         'for student field investigations, community monitoring walks, and '
         '"Be a Community Air Monitor" module (Section 7.1). Distributed across '
         '4 pilot communities (15 per community)',
         '$99', '60', '$5,940'],
        ['Subtotal B: Portable Sensors', '', '', '', '$5,940'],

        # --- Deployment & installation ---
        ['', 'C. DEPLOYMENT & INSTALLATION', '', '', ''],
        ['Site survey & installation labor',
         'Electrician/facilities support for mounting fixed sensors at school/community '
         'sites (avg. 2 hours per site)',
         '$150', '30', '$4,500'],
        ['WiFi access point (outdoor)',
         'Weatherproof WiFi extender for sites without outdoor WiFi coverage',
         '$80', '10', '$800'],
        ['Signage & QR code placards',
         'Weatherproof signs at each sensor site explaining the project and linking '
         'to real-time data on the BREATHE platform',
         '$25', '30', '$750'],
        ['Shipping & handling',
         'Inbound shipping of sensors, hardware, and supplies to 4 pilot regions',
         '$100', '4', '$400'],
        ['Subtotal C: Deployment & Installation', '', '', '', '$6,450'],

        # --- Calibration & QC ---
        ['', 'D. CALIBRATION & QUALITY CONTROL', '', '', ''],
        ['Co-location calibration kits',
         'Portable reference-grade PM2.5 sampler rental (2-week co-location '
         'calibration at each pilot region, 1x/year)',
         '$500', '4', '$2,000'],
        ['Calibration travel',
         'Travel to pilot sites for co-location setup and retrieval (2 trips/year \u00d7 4 regions)',
         '$350', '8', '$2,800'],
        ['Subtotal D: Calibration & QC', '', '', '', '$4,800'],

        # --- Maintenance ---
        ['', 'E. MAINTENANCE & REPLACEMENT (2 YEARS)', '', '', ''],
        ['Replacement laser counters',
         'PurpleAir replacement lasers for sensors showing drift (est. 20% failure '
         'rate over 2 years)',
         '$25', '12', '$300'],
        ['Replacement PurpleAir Flex sensors',
         'Full replacement units for damaged/failed sensors (est. 3 units over 2 years)',
         '$289', '3', '$867'],
        ['Replacement AirBeam Mini sensors',
         'Replacement for lost/damaged portable units (est. 10% attrition)',
         '$99', '6', '$594'],
        ['Sensor cleaning supplies',
         'Compressed air, microfiber cloths, isopropyl alcohol (per region/year)',
         '$25', '8', '$200'],
        ['Subtotal E: Maintenance & Replacement', '', '', '', '$1,961'],

        # --- Data & platform ---
        ['', 'F. DATA MANAGEMENT & PLATFORM INTEGRATION', '', '', ''],
        ['BREATHE platform sensor integration',
         'Development of data ingestion pipeline for PurpleAir API and AirCasting API '
         'into the BREATHE Data Exploration Platform (one-time development)',
         '$2,000', '1', '$2,000'],
        ['Cloud hosting for sensor data',
         'Cloud storage and compute for real-time sensor data processing, '
         'QC algorithms, and platform serving (Year 1\u20132)',
         '$100/mo', '24 mo', '$2,400'],
        ['Subtotal F: Data Management & Platform', '', '', '', '$4,400'],
    ],
    totals=['TOTAL STAGE 1', '', '', '', '$35,171'],
    note='Fixed sensor deployment: Houston (8 sites), Galveston (6 sites), El Paso (8 sites), '
         'Brownsville (8 sites) = 30 fixed sites. Portable sensors: 15 AirBeam Mini per pilot '
         'community = 60 units total. Prices based on PurpleAir (purpleair.com) and HabitatMap '
         'AirBeam (habitatmap.org) manufacturer list prices as of 2025\u201326.'
)

# ============================================================
# STAGE 2
# ============================================================
add_budget_table(
    'Stage 2: Observation-Based Citizen Science Across Texas (Years 2\u20134)',
    ['Item', 'Description', 'Unit Cost', 'Qty', 'Total'],
    [
        ['', 'A. THIRD-PARTY DATA INTEGRATION', '', '', ''],
        ['Third-party sensor data integration',
         'API development for importing PurpleAir community network and other '
         'third-party low-cost sensor data (e.g., OpenAQ) into BREATHE platform',
         '$2,500', '1', '$2,500'],
        ['Subtotal A: Third-Party Integration', '', '', '', '$2,500'],

        ['', 'B. DEMONSTRATION SENSOR KITS FOR EXPANSION SITES', '', '', ''],
        ['AirBeam Mini demonstration kits',
         'Portable PM2.5 sensors for educator workshops and PD events at expansion '
         'sites; loaned to trained educators for classroom use (2-week loan periods)',
         '$99', '12', '$1,188'],
        ['Kit shipping (round-trip)',
         'Shipping loaner kits to expansion sites across Texas and return',
         '$30', '24', '$720'],
        ['Subtotal B: Demonstration Kits', '', '', '', '$1,908'],

        ['', 'C. ONLINE TRAINING & MATERIALS', '', '', ''],
        ['Citizen science training videos',
         'Production of 4\u20136 short training videos: sky observation protocol, '
         'data submission, interpreting satellite comparisons, sensor basics',
         '$1,500', '1', '$1,500'],
        ['Printed observation field guides',
         'Laminated pocket guides for sky/visibility observation protocol '
         '(English/Spanish); distributed at PD workshops',
         '$3.50', '200', '$700'],
        ['Subtotal C: Training & Materials', '', '', '', '$2,200'],

        ['', 'D. STAGE 1 SENSOR NETWORK CONTINUATION', '', '', ''],
        ['Cloud hosting (Year 3\u20134)',
         'Continued hosting for pilot sensor data and expanded observation data',
         '$100/mo', '24 mo', '$2,400'],
        ['Pilot sensor maintenance (Year 3\u20134)',
         'Ongoing replacement parts and maintenance for Stage 1 fixed sensors',
         '$500/yr', '2 yr', '$1,000'],
        ['Subtotal D: Stage 1 Continuation', '', '', '', '$3,400'],
    ],
    totals=['TOTAL STAGE 2', '', '', '', '$10,008'],
    note='No new fixed sensor deployments. Expansion participation through the BREATHE platform '
         'observation tools, educator-facilitated activities, and third-party sensor integration. '
         'Demonstration AirBeam kits are circulated across expansion sites via a loan model.'
)

# ============================================================
# STAGE 3
# ============================================================
add_budget_table(
    'Stage 3: National Scaling Through Platform and Educator Networks (Years 4\u20135)',
    ['Item', 'Description', 'Unit Cost', 'Qty', 'Total'],
    [
        ['', 'A. PLATFORM OPERATIONS', '', '', ''],
        ['Cloud hosting (Year 4\u20135)',
         'Scaled hosting for national observation data, community map, '
         'and continued pilot sensor feeds',
         '$125/mo', '24 mo', '$3,000'],
        ['Platform maintenance & updates',
         'Bug fixes, new data layer integration, observation tool improvements',
         '$1,500/yr', '2 yr', '$3,000'],
        ['Subtotal A: Platform Operations', '', '', '', '$6,000'],

        ['', 'B. NATIONAL CITIZEN SCIENCE MATERIALS', '', '', ''],
        ['Observation protocol adaptation',
         'Adapting sky/visibility and environmental observation protocols for '
         'diverse climates and regions beyond Texas (review, testing, revision)',
         '$1,500', '1', '$1,500'],
        ['Printed/digital observation guides',
         'National distribution of observation field guides (English/Spanish) '
         'through educator networks and partner organizations',
         '$3.50', '500', '$1,750'],
        ['AirBeam Mini demonstration kits',
         'Additional portable kits for national PD workshops and partner orgs',
         '$99', '8', '$792'],
        ['Kit shipping (round-trip)',
         'Shipping loaner kits nationally',
         '$45', '16', '$720'],
        ['Subtotal B: National Materials', '', '', '', '$4,762'],

        ['', 'C. PILOT SENSOR NETWORK SUNSET', '', '', ''],
        ['Final calibration & data archival',
         'End-of-project co-location calibration, data quality assessment, '
         'and archival of full sensor datasets to public repository',
         '$500', '4', '$2,000'],
        ['Sensor decommissioning or transfer',
         'Transfer functioning sensors to partner schools/orgs for continued '
         'use; decommission non-functional units',
         '$50', '30', '$1,500'],
        ['Subtotal C: Sensor Sunset', '', '', '', '$3,500'],
    ],
    totals=['TOTAL STAGE 3', '', '', '', '$14,262'],
    note='No new fixed sensor deployments at national scale. All citizen science participation through '
         'the BREATHE platform observation tools, learning modules (Section 7.1), and educator-facilitated '
         'activities. Demonstration sensor kits circulated via loan model for workshops.'
)

# ============================================================
# GRAND TOTAL SUMMARY
# ============================================================
add_heading('Budget Summary: Section 7.4 Citizen Science and Community Monitoring', size=Pt(11))

summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = summary_table.rows[0]
for i, text in enumerate(['Project Stage', 'Period', 'Total Cost']):
    set_cell(hdr.cells[i], text, bold=True, size=Pt(10), align='center')
shade_row(hdr, '003366')
for cell in hdr.cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

summary_rows = [
    ['Stage 1: Sensor-Based Monitoring in Pilot Communities', 'Years 1\u20132', '$35,171'],
    ['Stage 2: Observation-Based Citizen Science Across Texas', 'Years 2\u20134', '$10,008'],
    ['Stage 3: National Scaling Through Platform & Educator Networks', 'Years 4\u20135', '$14,262'],
]
for row_data in summary_rows:
    row = summary_table.add_row()
    for i, val in enumerate(row_data):
        align = 'right' if i == 2 else 'left'
        set_cell(row.cells[i], val, size=Pt(10), align=align)

total_row = summary_table.add_row()
set_cell(total_row.cells[0], 'GRAND TOTAL (5 YEARS)', bold=True, size=Pt(10))
set_cell(total_row.cells[1], 'Years 1\u20135', bold=True, size=Pt(10), align='center')
set_cell(total_row.cells[2], '$59,441', bold=True, size=Pt(10), align='right')
shade_row(total_row, '003366')
for cell in total_row.cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

# Cost per category breakdown
add_heading('Cost Breakdown by Category', size=Pt(11))

cat_table = doc.add_table(rows=1, cols=3)
cat_table.style = 'Table Grid'
cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr2 = cat_table.rows[0]
for i, text in enumerate(['Category', '5-Year Total', '% of Budget']):
    set_cell(hdr2.cells[i], text, bold=True, size=Pt(10), align='center')
shade_row(hdr2, '003366')
for cell in hdr2.cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)

categories = [
    ['Fixed sensor hardware (PurpleAir Flex \u00d7 30)', '$11,620', '19.5%'],
    ['Portable sensor hardware (AirBeam Mini \u00d7 60)', '$5,940', '10.0%'],
    ['Deployment & installation', '$6,450', '10.9%'],
    ['Calibration & quality control', '$4,800', '8.1%'],
    ['Maintenance & replacement parts (all stages)', '$2,961', '5.0%'],
    ['Platform integration & third-party APIs', '$4,500', '7.6%'],
    ['Cloud hosting (5 years)', '$7,800', '13.1%'],
    ['Platform maintenance & updates (Stage 3)', '$3,000', '5.0%'],
    ['Training materials & videos', '$2,200', '3.7%'],
    ['Demonstration kit loans & shipping (Stages 2\u20133)', '$3,420', '5.8%'],
    ['Observation guides & protocol adaptation', '$3,250', '5.5%'],
    ['Sensor decommissioning & archival', '$3,500', '5.9%'],
]
for row_data in categories:
    row = cat_table.add_row()
    for i, val in enumerate(row_data):
        align = 'right' if i >= 1 else 'left'
        set_cell(row.cells[i], val, size=Pt(9), align=align)

total_cat = cat_table.add_row()
set_cell(total_cat.cells[0], 'TOTAL', bold=True, size=Pt(10))
set_cell(total_cat.cells[1], '$59,441', bold=True, size=Pt(10), align='right')
set_cell(total_cat.cells[2], '100%', bold=True, size=Pt(10), align='right')
shade_row(total_cat, 'E8EEF4')

# Notes
add_heading('Budget Notes and Assumptions', size=Pt(11))

notes = [
    '1. Sensor pricing: PurpleAir Flex $289, outdoor power supply $45, '
    'replacement laser $25, MicroSD $10 (purpleair.com, March 2026). '
    'AirBeam Mini $99 (habitatmap.org, March 2026).',

    '2. Fixed sensor allocation (30 units): Houston (8 sites: 4 schools, 2 community '
    'centers, 2 Ship Channel-adjacent), Galveston (6 sites: 3 schools, 2 coastal, '
    '1 community center), El Paso (8 sites: 4 schools, 2 border corridor, 2 desert/dust), '
    'Brownsville/RGV (8 sites: 4 schools, 2 agricultural areas, 2 community health centers).',

    '3. Portable sensor allocation (60 units): 15 AirBeam Mini per pilot community, '
    'supporting the "Be a Community Air Monitor" learning module (Section 7.1), '
    'student field investigations, and community monitoring walks.',

    '4. Calibration: Annual 2-week co-location with a reference-grade PM2.5 sampler '
    'at each pilot region to generate site-specific correction factors, consistent '
    'with EPA guidance on low-cost sensor use.',

    '5. Stage 2\u20133 scaling: No new fixed sensor networks. Participation scales through '
    '(a) sky/visibility observations on the BREATHE platform, (b) third-party '
    'sensor data integration (PurpleAir community network, OpenAQ), and (c) loaner '
    'AirBeam Mini demonstration kits circulated to educator workshops.',

    '6. Cloud hosting: Leverages university infrastructure where possible. Costs '
    'cover data ingestion from PurpleAir and AirCasting APIs, observation storage, '
    'QC processing, and serving data layers to the BREATHE Data Exploration Platform '
    '(Section 7.2). Scales from $100/mo (Stage 1) to $125/mo (Stage 3, national).',

    '7. All sensor hardware and materials costs exclude applicable sales tax, which '
    'may be exempt under university institutional purchasing agreements.',
]

for note in notes:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(note)
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'

# Save
output_path = r'd:\_NASA_F9_Breath\Claude_project\BREATHE_Section_7_4_Budget.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
