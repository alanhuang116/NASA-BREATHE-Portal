"""Generate Section 7.4 for the BREATHE proposal as a Word document."""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from lxml import etree
import re

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.25


def add_breathe_header():
    header_para = doc.add_paragraph()
    parts = [
        ('BREATHE: ', True, False),
        ('BR', True, True),
        ('idging ', True, False),
        ('E', True, True),
        ('ducation on ', True, False),
        ('A', True, True),
        ('tmosphere, ', True, False),
        ('T', True, True),
        ('echnology, and ', True, False),
        ('HE', True, True),
        ('alth', True, False),
    ]
    for text, bold, underline in parts:
        r = header_para.add_run(text)
        r.bold = bold
        r.underline = underline
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'

    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_after = Pt(6)
    pPr = border_para._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')


def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    return p


def add_body(text, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.25
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'
    return p


# ============================================================
# SECTION 7.4 — Citizen Science and Community Monitoring
# ============================================================
add_breathe_header()
add_section_heading('7.4 Citizen Science and Community Monitoring')

# Paragraph 1: Framing, theory, and tiered design
add_body(
    'BREATHE engages participants in collecting environmental measurements and comparing '
    'them with NASA satellite observations, following a tiered implementation strategy that '
    'matches the depth of citizen science activity to available resources at each project '
    'stage. This approach is grounded in the framework of **legitimate peripheral participation** '
    '[Lave and Wenger, 1991], which holds that learners develop scientific competence by '
    'engaging in authentic practices within a community of practice\u2014beginning with structured, '
    'guided tasks and progressively taking on more independent roles as their expertise grows. '
    'In BREATHE, participants move from guided observation and data interpretation to '
    'independent sensor deployment and ground-truth validation, building scientific identity '
    'through increasingly consequential contributions to NASA Earth science.'
)

# Paragraph 2: Stage 1 — sensor deployment in pilot communities
add_body(
    '**Stage 1: Sensor-Based Monitoring in Pilot Communities.** During the first two years, '
    'BREATHE deploys networks of **low-cost air quality sensors** (PM2.5 monitors) at partner '
    'schools and community sites in Houston, Galveston, El Paso, and Brownsville. Working '
    'through the learning module \u201cBe a Community Air Monitor\u201d (Section 7.1), student and '
    'community participants assemble, calibrate, and deploy sensors; design sampling strategies '
    'appropriate to their local environment (e.g., proximity to the Houston Ship Channel, '
    'agricultural burning cycles in the Rio Grande Valley, transboundary dust events in El '
    'Paso); and collect continuous PM2.5 time-series data. Participants follow standardized '
    'data collection and quality-control protocols developed in consultation with HAQAST '
    'scientists, ensuring that measurements meet minimum quality thresholds for meaningful '
    'comparison with satellite products. Sensor data are uploaded to the BREATHE Data '
    'Exploration Platform (Section 7.2), where they appear alongside TEMPO, MAIA, and '
    'MODIS/VIIRS satellite layers, enabling participants to visually compare their ground-level '
    'readings with satellite-derived estimates for the same location and time period. This '
    'co-visualization makes the concept of satellite validation concrete: participants can see '
    'where satellite and ground measurements agree, where they diverge, and discuss what '
    'factors\u2014cloud cover, sensor placement, spatial resolution\u2014might explain differences. '
    'We target deployment of sensors at approximately 10\u201315 sites across the four pilot '
    'communities during Stage 1, with each site maintained by a trained educator or community '
    'partner.'
)

# Paragraph 3: Stage 2 — Texas expansion without full sensor networks
add_body(
    '**Stage 2: Observation-Based Citizen Science Across Texas.** As BREATHE expands across '
    'Texas (years 2\u20134), deploying dedicated sensor networks at every new site is not feasible. '
    'Instead, the project shifts to a **virtual and observation-based citizen science model** '
    'that enables broad participation without hardware costs. Participants at expansion sites '
    'engage through three complementary pathways. First, they use the Data Exploration Platform '
    'to conduct **guided satellite data investigations** for their own communities\u2014selecting '
    'their location, examining TEMPO and MAIA data layers, comparing current conditions with '
    'historical patterns, and recording structured observations about local environmental '
    'conditions (e.g., visible haze, nearby emission sources, traffic patterns, weather) that '
    'provide qualitative ground context for satellite data. Second, participants submit '
    '**standardized sky and visibility observations** through a mobile-friendly interface on '
    'the platform, reporting cloud cover, sky clarity, and perceived air quality using a '
    'structured protocol adapted from GLOBE\u2019s atmosphere observation methods. These '
    'observations are timestamped and geolocated, creating a distributed network of human '
    'observers whose reports complement satellite measurements\u2014particularly valuable in areas '
    'with limited ground-based monitoring infrastructure, such as rural and border communities. '
    'Third, where participants or partner organizations have access to existing low-cost sensor '
    'networks (e.g., PurpleAir), the platform provides tools for **importing and co-visualizing '
    'third-party sensor data** alongside NASA products, extending the ground-to-satellite '
    'comparison without requiring BREATHE-funded hardware. Online training modules (Section '
    '7.1) prepare participants for these activities with the same data literacy skills taught '
    'in the sensor-based pilot, ensuring scientific rigor is maintained as participation scales.'
)

# Paragraph 4: Stage 3 — National scaling through the platform
add_body(
    '**Stage 3: National Scaling Through the Platform and Educator Networks.** At the national '
    'scale (years 4\u20135), citizen science participation is sustained entirely through the Data '
    'Exploration Platform and online learning modules, with no project-funded sensor deployment. '
    'Participants nationwide use the platform\u2019s place-based interface to investigate satellite-'
    'observed air quality in their communities, submit sky and environmental observations, '
    'and compare their reports with TEMPO, PACE, and MAIA data. The platform aggregates these '
    'distributed observations into a **community observation map** that grows as participation '
    'expands, providing a nationally scaled, participant-generated dataset of local atmospheric '
    'conditions that contextualizes NASA satellite products. Educators trained through '
    'BREATHE\u2019s professional development pathway (Section 7.3) facilitate these activities in '
    'classrooms and informal learning settings using the learning modules and data activities '
    'from Section 7.1. Where sensor-based pilot data from Stage 1 communities remain active, '
    'they continue to feed the platform and serve as anchoring examples for new participants\u2014'
    'demonstrating what ground-truth validation looks like and motivating deeper engagement. '
    'This tiered design ensures that BREATHE\u2019s citizen science activities are scientifically '
    'meaningful at every scale: hardware-intensive where resources allow, observation-based '
    'where they do not, and always connected to authentic NASA satellite data through the '
    'platform.'
)

# Save
output_path = r'd:\_NASA_F9_Breath\Claude_project\BREATHE_Section_7_4.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
