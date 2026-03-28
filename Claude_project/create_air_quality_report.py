"""Generate a detailed MS Word report on NASA Air Quality Education Efforts."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table_row(table, cells_data, bold_first=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if bold_first and i == 0:
            run.bold = True
    return row

def shade_cells(row, color):
    for cell in row.cells:
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): color,
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NASA Air Quality Education & Citizen Science Efforts')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('A Comprehensive Summary of Programs Across\nGLOBE Mission Earth, My NASA Data, and NESEC')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'Prepared: {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. GLOBE Program — Mission Earth & Atmosphere Investigations',
    '   2.1 Overview of Mission Earth',
    '   2.2 GLOBE Atmosphere Protocols (All Nine)',
    '   2.3 Aerosol Protocol — Core Air Quality Measurement',
    '   2.4 Clouds Protocol — Satellite Validation',
    '   2.5 Supporting Protocols for Air Quality Context',
    '   2.6 GLOBE Observer App',
    '   2.7 Campaigns & Intensive Observation Periods',
    '3. My NASA Data (mynasadata.larc.nasa.gov)',
    '   3.1 Platform Overview',
    '   3.2 Air Quality Mini-Lessons & Activities',
    '   3.3 Air Quality Lesson Plans (Archived)',
    '   3.4 Interactive Tools & Data Explorer',
    '   3.5 NASA Missions & Data Sources',
    '4. NASA Earth Science Education Collaborative (NESEC)',
    '   4.1 Mission & Role',
    '   4.2 Resource Portfolio',
    '5. Cross-Cutting Themes & Synergies',
    '6. Key NASA Satellite Missions Supporting Air Quality Education',
    '7. Conclusions',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading_styled('1. Executive Summary', level=1)

doc.add_paragraph(
    'NASA maintains a robust ecosystem of air quality education and citizen science programs '
    'that connect students, educators, and the public to authentic atmospheric research. This report '
    'provides a detailed examination of air quality-related efforts across three major NASA-supported platforms:'
)

add_bullet('GLOBE Program (globe.gov) — Mission Earth & Atmosphere Investigations', bold_prefix=None)
add_bullet('My NASA Data (mynasadata.larc.nasa.gov) — NASA Langley Research Center', bold_prefix=None)
add_bullet('NASA Earth Science Education Collaborative (nesec.strategies.org)', bold_prefix=None)

doc.add_paragraph(
    'Together, these programs enable participants worldwide to collect ground-based atmospheric measurements '
    'using standardized protocols, validate NASA satellite observations, analyze authentic air quality datasets, '
    'and develop scientific literacy about the causes and consequences of air pollution. The programs span grades '
    '3–12 and extend to adult citizen scientists, operating across Africa, Asia-Pacific, Europe, Latin America, '
    'and North America.'
)

doc.add_paragraph(
    'Key air quality topics addressed include aerosol optical thickness, particulate matter (PM2.5), '
    'carbon monoxide pollution, ozone monitoring, volcanic ash and aerosol transport, wildfire smoke modeling, '
    'acid precipitation pH, and the health impacts of air pollution. These educational efforts are directly '
    'tied to NASA missions including CALIPSO, Terra, Aqua (MODIS), CERES, and the broader Earth Observing System.'
)

# ============================================================
# 2. GLOBE PROGRAM
# ============================================================
add_heading_styled('2. GLOBE Program — Mission Earth & Atmosphere Investigations', level=1)

# 2.1
add_heading_styled('2.1 Overview of Mission Earth', level=2)
doc.add_paragraph(
    'Mission Earth is the overarching initiative that fuses GLOBE\'s global citizen science network with '
    'NASA\'s satellite and research assets to drive systematic innovation in STEM education. The program '
    'engages participants in hands-on environmental data collection across five Earth system domains: '
    'Atmosphere, Biosphere, Hydrosphere, Pedosphere (soil science), and integrated Earth-as-a-System bundles.'
)
doc.add_paragraph(
    'Within the atmospheric domain, Mission Earth connects students to structured measurement protocols that '
    'generate scientifically rigorous data used by researchers worldwide. Participants contribute to Current '
    'Global Measurement Campaigns and Intensive Observation Periods (IOPs), enabling coordinated data collection '
    'events that coincide with satellite overpasses and seasonal phenomena.'
)
doc.add_paragraph(
    'Mission Earth operates through multiple access points including the GLOBE Observer mobile app for field '
    'data collection, desktop data entry systems, regional collaboration frameworks, and student research '
    'exhibition opportunities. The program maintains an open data set available to scientists globally.'
)

# 2.2
add_heading_styled('2.2 GLOBE Atmosphere Protocols — Complete Overview', level=2)
doc.add_paragraph(
    'GLOBE offers nine core atmospheric investigation protocols that collectively provide a comprehensive '
    'framework for understanding weather, climate, and atmospheric composition — the three major research '
    'focus areas. These protocols are designed for citizen scientists of all ages and are accompanied by '
    'field guides, data sheets, learning activities, and classroom preparation materials.'
)

# Protocols table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0]
for i, text in enumerate(['Protocol', 'Measurement Focus', 'Instruments', 'Air Quality Relevance']):
    hdr.cells[i].text = ''
    p = hdr.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
shade_cells(hdr, '003366')

protocols = [
    ['Aerosols', 'Aerosol Optical Thickness (AOT)', 'GLOBE Sun Photometer + digital voltmeter',
     'DIRECT — Primary air quality protocol; measures airborne particle concentration'],
    ['Air Temperature', 'Current air temperature (shade, 3+ min)', 'Digital thermometers, automated weather stations (Davis, RainWise, WeatherHawk), HOBO loggers',
     'SUPPORTING — Temperature affects pollutant dispersion, atmospheric stability, and chemical reaction rates'],
    ['Barometric Pressure', 'Atmospheric pressure', 'Aneroid barometer',
     'SUPPORTING — Required for Aerosols and Water Vapor protocols; pressure systems drive pollutant transport'],
    ['Clouds', 'Cloud type, sky coverage, opacity', 'Visual observation + field guides',
     'DIRECT — Cloud observations validate NASA satellite data; clouds affect pollutant washout and photochemistry'],
    ['Precipitation', 'Rainfall/snowfall amount + pH', 'Rain gauges, snow boards, pH meters',
     'DIRECT — pH measurement tracks acid precipitation from atmospheric pollution'],
    ['Relative Humidity', 'Atmospheric moisture levels', 'Digital hygrometer or sling psychrometer',
     'SUPPORTING — Humidity affects aerosol behavior, particle size, and smog formation'],
    ['Surface Temperature', 'Earth surface temperature', 'Infrared thermometer (IRT)',
     'SUPPORTING — Surface temperature drives convection and pollutant mixing in the boundary layer'],
    ['Water Vapor', 'Atmospheric water vapor content', 'Near-infrared sun photometer',
     'SUPPORTING — Water vapor affects atmospheric transparency and aerosol interactions'],
    ['Wind', 'Wind direction (+ ozone stand construction)', 'Constructed wind direction instrument + ozone measurement stand',
     'DIRECT — Wind direction determines pollutant transport; protocol includes ozone measurement stand construction'],
]

for row_data in protocols:
    r = add_table_row(table, row_data, bold_first=True)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(4.0)
    row.cells[2].width = Cm(4.5)
    row.cells[3].width = Cm(6.0)

# 2.3
add_heading_styled('2.3 Aerosol Protocol — Core Air Quality Measurement', level=2)
doc.add_paragraph(
    'The GLOBE Aerosol Protocol is the program\'s most directly air-quality-relevant measurement activity. '
    'It measures Aerosol Optical Thickness (AOT) — a quantitative measure of how much sunlight is scattered '
    'or absorbed by suspended particles (dust, smoke, industrial pollution, pollen, sea salt) in the atmosphere.'
)

add_heading_styled('Measurement Procedure', level=3)
doc.add_paragraph('Students and citizen scientists follow a structured multi-step observation process:')
add_bullet('Point the GLOBE Sun Photometer directly at the sun')
add_bullet('Record the maximum (largest) voltage reading on a connected digital voltmeter')
add_bullet('Observe and document sky conditions near the sun (haze, clarity)')
add_bullet('Record the exact time of measurement for satellite data matching')
add_bullet('Conduct all required complementary measurements simultaneously')

add_heading_styled('Required Supporting Observations', level=3)
doc.add_paragraph(
    'Each aerosol measurement session requires coordination with several complementary protocols to ensure '
    'comprehensive atmospheric characterization:'
)
add_bullet('Barometric Pressure: ', bold_prefix=None)
add_bullet('Cloud type, sky coverage percentage, and cloud opacity', bold_prefix=None)
add_bullet('Current air temperature', bold_prefix=None)
add_bullet('Relative humidity', bold_prefix=None)

add_heading_styled('NASA Satellite Validation', level=3)
doc.add_paragraph(
    'Ground-based aerosol measurements are a critical component of NASA\'s satellite validation program. '
    'Student observations are timed to coincide with NASA satellite overpasses. Participants access satellite '
    'overpass schedules through NASA\'s Cloud Satellite Portal to optimize their observation timing. '
    'The combination of ground-level ("looking up") and satellite ("looking down") perspectives provides '
    'scientists with a much more complete picture of atmospheric aerosol conditions than either vantage point alone.'
)

# 2.4
add_heading_styled('2.4 Clouds Protocol — Satellite Validation for Atmospheric Science', level=2)
doc.add_paragraph(
    'The GLOBE Clouds Protocol enables citizen scientists to observe and classify cloud types, estimate '
    'sky coverage percentage, and assess cloud opacity. This protocol is particularly significant for air '
    'quality science because clouds interact with atmospheric pollutants through washout processes, affect '
    'photochemical reaction rates, and influence the Earth\'s energy budget.'
)

add_heading_styled('Key Features', level=3)
add_bullet('Flexible timing — designed to fit observers\' schedules', bold_prefix=None)
add_bullet('Each observation is matched to satellite data collected at approximately the same time and location', bold_prefix=None)
add_bullet('Participants receive email comparisons between their ground observations and satellite data', bold_prefix=None)
add_bullet('Dual-perspective approach: "Satellites only see the top of the clouds while you see the bottom. '
           'By putting these two vantage points together we get a much more complete picture."', bold_prefix=None)

add_heading_styled('Data Collection Elements', level=3)
add_bullet('Cloud type classification (using standardized field guides)', bold_prefix=None)
add_bullet('Percentage of sky covered by clouds', bold_prefix=None)
add_bullet('Cloud opacity (transparency/thickness)', bold_prefix=None)
add_bullet('Sky and surface conditions at time of observation', bold_prefix=None)

# 2.5
add_heading_styled('2.5 Supporting Protocols for Air Quality Context', level=2)

doc.add_paragraph(
    'Several additional GLOBE atmosphere protocols provide essential contextual data that enhances '
    'the interpretation of air quality measurements:'
)

add_heading_styled('Precipitation Protocol — Acid Rain Detection', level=3)
doc.add_paragraph(
    'Beyond measuring rainfall and snowfall amounts, this protocol includes pH measurement of collected '
    'precipitation. This directly monitors acid precipitation — a key indicator of atmospheric pollution '
    'from sulfur dioxide (SO₂) and nitrogen oxides (NOₓ) emissions. Students use rain gauges, snow boards, '
    'and pH analysis equipment (various methods and materials) with standardized lab guides.'
)

add_heading_styled('Wind Protocol — Pollutant Transport', level=3)
doc.add_paragraph(
    'The Wind protocol includes the "Ozone Measurement Stand and Wind Direction Instrument Construction '
    'Protocol," which provides instructions for building both a wind direction instrument and an ozone '
    'measurement station. Wind data is essential for understanding how atmospheric circulation patterns '
    'distribute air pollutants and influence regional and global air quality conditions. The protocol '
    'integrates with automated weather stations (Davis, RainWise, WeatherHawk, Earth Networks) that collect '
    'multiple atmospheric parameters at 15-minute intervals.'
)

add_heading_styled('Relative Humidity Protocol', level=3)
doc.add_paragraph(
    'Students measure atmospheric moisture using digital hygrometers or sling psychrometers. Humidity '
    'directly affects aerosol behavior — high humidity causes hygroscopic particles to swell, changing '
    'their optical properties and health impacts. Humidity also influences smog formation rates and the '
    'conversion of gaseous pollutants to particulate matter.'
)

add_heading_styled('Water Vapor Protocol', level=3)
doc.add_paragraph(
    'Using a near-infrared sun photometer, students measure sunlight at wavelengths correlated to atmospheric '
    'water vapor content. This measurement complements aerosol optical thickness data by helping scientists '
    'distinguish between water vapor absorption and particle scattering in the atmosphere — essential for '
    'accurate air quality assessment from both ground and satellite instruments.'
)

add_heading_styled('Air Temperature & Surface Temperature Protocols', level=3)
doc.add_paragraph(
    'Temperature data (both ambient air and surface) is fundamental to understanding atmospheric stability '
    'and boundary layer dynamics. Temperature inversions trap pollutants near the surface, creating poor '
    'air quality episodes. Students measure air temperature with digital thermometers held in shade for '
    'at least 3 minutes, and surface temperature with infrared thermometers (IRTs). Automated weather '
    'stations and HOBO data loggers provide continuous 15-minute interval recordings.'
)

# 2.6
add_heading_styled('2.6 GLOBE Observer App', level=2)
doc.add_paragraph(
    'GLOBE Observer is a free citizen science mobile application that allows volunteers in GLOBE-participating '
    'countries to take environmental observations and contribute to the global GLOBE community. The app '
    'supports four primary observation types:'
)
add_bullet('Cloud Cover — sky condition measurements that support atmospheric and air quality science', bold_prefix=None)
add_bullet('Land Cover Classification — surface conditions affecting air quality (urban heat islands, vegetation)', bold_prefix=None)
add_bullet('Mosquito Habitats — environmental health monitoring', bold_prefix=None)
add_bullet('Tree Height — canopy measurements relevant to air filtration and carbon sequestration', bold_prefix=None)

doc.add_paragraph(
    'The app enables rapid field data collection that feeds into GLOBE\'s open database, supporting '
    'both NASA\'s Earth system science research and educational investigations. Cloud observations '
    'submitted through the app are matched with satellite data for validation purposes.'
)

# 2.7
add_heading_styled('2.7 Campaigns & Intensive Observation Periods', level=2)
doc.add_paragraph(
    'GLOBE organizes Current Global Measurement Campaigns and Intensive Observation Periods (IOPs) that '
    'coordinate data collection across its worldwide network spanning Africa, Asia-Pacific, Europe, '
    'Latin America, and North America. These campaigns often focus on specific atmospheric phenomena '
    'and are timed to coincide with satellite mission activities, seasonal events (e.g., wildfire season, '
    'dust storm periods), or specific research questions.'
)
doc.add_paragraph(
    'Student participants can showcase their atmospheric research findings through GLOBE\'s Student '
    'Research Exhibition program, providing opportunities to present original air quality investigations '
    'to the scientific community.'
)

# ============================================================
# 3. MY NASA DATA
# ============================================================
add_heading_styled('3. My NASA Data (mynasadata.larc.nasa.gov)', level=1)

# 3.1
add_heading_styled('3.1 Platform Overview', level=2)
doc.add_paragraph(
    'My NASA Data is a NASA Langley Research Center Science Directorate project that provides authentic '
    'NASA Earth data for educators and learners in grades 3–12. The platform curates NASA datasets and '
    'packages them into standards-aligned educational resources organized by Earth system phenomena. '
    '"Air Quality" is a dedicated phenomenon topic within the Atmosphere section.'
)
doc.add_paragraph(
    'The platform defines air quality as "a measure of the pollution level in the air" and emphasizes '
    'that "monitoring air quality is important because polluted air can be bad for human health and '
    'the environment." Resources are available in multiple formats including Mini-Lessons, Interactives, '
    'Lesson Plans, and Data Resources.'
)

# 3.2
add_heading_styled('3.2 Air Quality Mini-Lessons & Activities', level=2)
doc.add_paragraph(
    'My NASA Data provides a curated collection of air quality mini-lessons that use authentic NASA and '
    'EPA data to teach atmospheric science concepts. Each mini-lesson is designed for specific grade bands '
    'and includes structured activities with Google Docs/Slides integration.'
)

# Mini-lessons table
table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0]
for i, text in enumerate(['Mini-Lesson Title', 'Grade Level', 'Description', 'Data Sources']):
    hdr2.cells[i].text = ''
    p = hdr2.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
shade_cells(hdr2, '003366')

mini_lessons = [
    ['How is My Air?', '6–8',
     'Students use AirNow.gov to assess current real-time air quality conditions at US locations, interpret Air Quality Index (AQI) values, and explore factors affecting local air quality.',
     'AirNow.gov (EPA)'],
    ['Identifying Patterns in PM 2.5', '9–12',
     'Students analyze global maps showing population distribution, PM2.5 (fine particulate matter) concentrations, and pollution-attributable mortality. Examines the quantitative relationship between particulate air pollution and human health outcomes across countries and regions.',
     'NASA Earth Observatory, Global health datasets'],
    ['Multiyear Time Plots for Air Quality Data', '6–8',
     'Learners analyze temporal trends in air quality data, tracking how pollution levels have changed over multiple years at selected US locations. Develops data literacy and trend analysis skills.',
     'EPA AirData Portal'],
    ['Energy and Matter: Dust Transport', '6–8, 9–12',
     'Features NASA CALIPSO satellite video observations showing trans-Atlantic dust migration from the Sahara Desert to the Amazon rainforest. Explores how natural aerosol transport connects distant ecosystems and affects air quality across continents.',
     'NASA CALIPSO satellite mission'],
    ['Kuril Islands Volcanoes', '3–5, 6–8',
     'Compares satellite imagery of volcanic eruptions from two separate events occurring ten years apart. Students analyze how volcanic emissions affect multiple Earth spheres, including atmospheric aerosol loading.',
     'NASA Earth Observatory satellite imagery'],
    ['Analyzing a Volcanic Ash Model', '6–8, 9–12',
     'Explores volcanic ash behavior through visualization of the Calbuco volcano eruption, demonstrating how volcanic aerosols disperse globally through atmospheric circulation.',
     'NASA volcanic ash dispersion models'],
    ['Introduction to Volcanic Ash', '6–8, 9–12',
     'Presents NASA videos addressing aerosols and volcanic ash phenomena, introducing students to the concept of atmospheric particulate matter from volcanic sources.',
     'NASA video resources'],
    ['Systems and System Models: Observing Our Planet on Fire', '6–8',
     'Uses fire smoke modeling to illustrate Earth system interconnections and predictive forecasting. Shows how wildfire smoke affects air quality across large regions.',
     'NASA fire and smoke models'],
]

for row_data in mini_lessons:
    r = add_table_row(table2, row_data, bold_first=True)

for row in table2.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(1.5)
    row.cells[2].width = Cm(8.0)
    row.cells[3].width = Cm(4.0)

# 3.3
add_heading_styled('3.3 Air Quality Lesson Plans (Archived Collection)', level=2)
doc.add_paragraph(
    'My NASA Data maintains an archive of air quality lesson plans, many dating back to 2004 and periodically '
    'updated. These provide more in-depth, multi-session investigations compared to mini-lessons. '
    'Key archived air quality lesson plans include:'
)

add_bullet('"Do Fireworks Create Air Pollution?" — Investigates the air quality impacts of fireworks displays '
           'using real monitoring data (updated 2025)', bold_prefix=None)
add_bullet('"Carbon Monoxide and Population Density" — Examines the spatial relationship between CO concentrations '
           'and human population centers using satellite data (2018)', bold_prefix=None)
add_bullet('"Monitoring Ozone in National Parks" — Analyzes ground-level ozone data from national park monitoring '
           'stations, connecting air quality to environmental protection (updated 2025)', bold_prefix=None)
add_bullet('"Monitoring Ozone in the Great Lakes Region" — Regional ozone analysis focusing on the Great Lakes '
           'airshed and its unique atmospheric chemistry (updated 2025)', bold_prefix=None)

# 3.4
add_heading_styled('3.4 Interactive Tools & Earth System Data Explorer', level=2)
doc.add_paragraph(
    'My NASA Data provides the Earth System Data Explorer, an interactive visualization tool built on '
    'Google Earth Engine that allows students to explore air quality and other Earth science datasets '
    'through a user-friendly interface. The platform organizes interactive content under the Atmosphere '
    'section, enabling students to manipulate and visualize atmospheric data without requiring programming '
    'skills or specialized software.'
)
doc.add_paragraph(
    'These interactives support data literacy by allowing students to create custom visualizations, '
    'compare datasets across time periods and geographic regions, and investigate cause-and-effect '
    'relationships between human activities and air quality outcomes.'
)

# 3.5
add_heading_styled('3.5 NASA Missions & Data Sources', level=2)
doc.add_paragraph(
    'My NASA Data\'s air quality resources draw upon data from multiple NASA missions and external partners:'
)

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0]
for i, text in enumerate(['Data Source', 'Type', 'Air Quality Application']):
    hdr3.cells[i].text = ''
    p = hdr3.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
shade_cells(hdr3, '003366')

sources = [
    ['NASA CALIPSO', 'Satellite mission (lidar)', 'Vertical profiles of aerosols and clouds; dust transport visualization'],
    ['NASA Earth Observatory', 'Satellite imagery & visualizations', 'Volcanic eruption imagery, aerosol plumes, fire/smoke events'],
    ['NASA Fire/Smoke Models', 'Predictive models', 'Wildfire smoke dispersion and air quality forecasting'],
    ['NASA Volcanic Ash Models', 'Dispersion models', 'Global volcanic ash transport and atmospheric aerosol loading'],
    ['EPA AirData Portal', 'Ground monitoring network', 'Historical US air quality data for trend analysis'],
    ['AirNow.gov (EPA)', 'Real-time monitoring', 'Current Air Quality Index (AQI) for US locations'],
]

for row_data in sources:
    add_table_row(table3, row_data, bold_first=True)

# ============================================================
# 4. NESEC
# ============================================================
add_heading_styled('4. NASA Earth Science Education Collaborative (NESEC)', level=1)

# 4.1
add_heading_styled('4.1 Mission & Role', level=2)
doc.add_paragraph(
    'The NASA Earth Science Education Collaborative (NESEC), hosted at nesec.strategies.org, serves as '
    'a curated clearinghouse that connects educators to NASA Earth science education resources across '
    'NASA\'s entire ecosystem. NESEC aggregates and organizes educational materials from multiple NASA '
    'centers, programs, and missions — including those from GLOBE, My NASA Data, and other sources — '
    'into a searchable resource portfolio.'
)
doc.add_paragraph(
    'NESEC\'s role in air quality education is primarily as a connector and amplifier: it ensures that '
    'air quality and atmospheric science resources developed across NASA\'s distributed network are '
    'discoverable, accessible, and aligned with educational standards. The collaborative helps educators '
    'find the right resources for their grade level, subject area, and instructional context.'
)

# 4.2
add_heading_styled('4.2 Resource Portfolio', level=2)
doc.add_paragraph(
    'NESEC\'s resource portfolio covers multiple Earth system topics, with atmosphere and air quality '
    'being a key category. The portfolio includes resources that can be filtered by:'
)
add_bullet('Topic (e.g., atmosphere, climate, weather, air quality)', bold_prefix=None)
add_bullet('Grade level and educational standards', bold_prefix=None)
add_bullet('Resource type (lesson plans, activities, data tools, multimedia)', bold_prefix=None)
add_bullet('NASA mission or program affiliation', bold_prefix=None)

doc.add_paragraph(
    'By serving as a central discovery point, NESEC reduces the fragmentation that can occur across '
    'NASA\'s numerous education initiatives and helps educators efficiently locate atmospheric science '
    'and air quality materials suited to their specific needs.'
)

# ============================================================
# 5. CROSS-CUTTING THEMES
# ============================================================
add_heading_styled('5. Cross-Cutting Themes & Synergies', level=1)

doc.add_paragraph(
    'Across all three platforms, NASA\'s air quality education efforts share several important common themes '
    'that reinforce their collective impact:'
)

add_heading_styled('5.1 Citizen Science Meets Satellite Validation', level=2)
doc.add_paragraph(
    'Ground-based observations collected through GLOBE protocols (especially aerosol and cloud data) are '
    'directly used to validate NASA satellite measurements. This gives students and citizen scientists a '
    'meaningful role in real scientific research — their data has genuine scientific value, not just '
    'educational value. The dual-perspective approach (ground "looking up" + satellite "looking down") '
    'produces more complete atmospheric characterizations than either vantage point alone.'
)

add_heading_styled('5.2 Authentic Data Literacy', level=2)
doc.add_paragraph(
    'All three programs emphasize working with authentic NASA and EPA datasets rather than textbook '
    'examples or simulated data. Students learn to interpret real Air Quality Index values (AirNow.gov), '
    'analyze actual PM2.5 concentration maps, track genuine multi-year air quality trends (EPA AirData), '
    'and examine real satellite imagery of aerosol events. This authentic data approach builds scientific '
    'literacy and critical thinking skills.'
)

add_heading_styled('5.3 Health-Environment Connections', level=2)
doc.add_paragraph(
    'Particularly through My NASA Data\'s PM2.5 lessons, the direct link between air pollution and human '
    'health is explicitly taught. Students analyze global data on pollution-attributable mortality, making '
    'the abstract concept of "air quality" tangible and personally relevant. This health-centered approach '
    'motivates engagement and connects atmospheric science to public health outcomes.'
)

add_heading_styled('5.4 Global-to-Local Scale Integration', level=2)
doc.add_paragraph(
    'Resources range from planetary-scale phenomena (trans-Atlantic dust transport via CALIPSO, global volcanic '
    'ash dispersion) to local-scale investigations (checking your community\'s AQI on AirNow.gov, measuring '
    'aerosol optical thickness at your school\'s study site). This multi-scale approach helps students '
    'understand how global atmospheric processes directly affect their own community\'s air quality.'
)

add_heading_styled('5.5 Integrated Earth System Science', level=2)
doc.add_paragraph(
    'Air quality is consistently presented not as an isolated topic but as part of interconnected Earth '
    'systems. Wildfire lessons connect biosphere burning to atmospheric smoke. Dust transport lessons link '
    'African deserts to Amazonian ecosystems. Volcanic eruption lessons show lithosphere events affecting '
    'atmospheric composition globally. Precipitation pH measurements connect atmospheric pollution to '
    'hydrosphere chemistry. This systems thinking approach reflects NASA\'s holistic Earth science framework.'
)

# ============================================================
# 6. KEY NASA MISSIONS
# ============================================================
add_heading_styled('6. Key NASA Satellite Missions Supporting Air Quality Education', level=1)

doc.add_paragraph(
    'The following NASA missions and instruments provide the foundational data that powers air quality '
    'education across all three platforms:'
)

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table4.rows[0]
for i, text in enumerate(['Mission / Instrument', 'Capability', 'Educational Use']):
    hdr4.cells[i].text = ''
    p = hdr4.cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(255, 255, 255)
shade_cells(hdr4, '003366')

missions = [
    ['CALIPSO (Cloud-Aerosol Lidar)', 'Lidar-based vertical profiles of aerosols and clouds',
     'Dust transport mini-lessons; aerosol vertical distribution visualization'],
    ['Terra / Aqua (MODIS)', 'Global aerosol optical depth, fire detection, true-color imagery',
     'Wildfire smoke lessons; volcanic plume imagery; global aerosol maps'],
    ['CERES', 'Earth\'s radiation budget and cloud properties',
     'Energy budget understanding; cloud-atmosphere interactions'],
    ['Earth Observing System (EOS)', 'Integrated Earth observation platform',
     'Foundation for all satellite-derived air quality educational datasets'],
    ['Cloud Satellite Portal', 'Satellite overpass scheduling tool',
     'GLOBE participants time ground observations to satellite passes for data validation'],
]

for row_data in missions:
    add_table_row(table4, row_data, bold_first=True)

# ============================================================
# 7. CONCLUSIONS
# ============================================================
add_heading_styled('7. Conclusions', level=1)

doc.add_paragraph(
    'NASA\'s air quality education ecosystem — spanning GLOBE Mission Earth, My NASA Data, and NESEC — '
    'represents a comprehensive, multi-faceted approach to atmospheric science education and citizen science. '
    'The programs collectively achieve several important outcomes:'
)

add_bullet('Enable global citizen science participation in authentic air quality data collection that '
           'directly supports NASA satellite validation', bold_prefix=None)
add_bullet('Provide standards-aligned, grade-appropriate educational resources (grades 3–12) that use '
           'real atmospheric data from NASA missions and EPA monitoring networks', bold_prefix=None)
add_bullet('Cover the full spectrum of air quality topics — from aerosol optical thickness measurement '
           'to PM2.5 health impacts, ozone monitoring, acid precipitation, volcanic aerosols, wildfire '
           'smoke, and dust transport', bold_prefix=None)
add_bullet('Connect local observations to global atmospheric processes through multi-scale investigations '
           'that build systems thinking', bold_prefix=None)
add_bullet('Maintain open, accessible databases and interactive tools that democratize access to NASA '
           'Earth science data', bold_prefix=None)
add_bullet('Operate across a global network spanning five continents, enabling coordinated international '
           'atmospheric monitoring campaigns', bold_prefix=None)

doc.add_paragraph(
    'These programs exemplify NASA\'s commitment to making Earth science research accessible and actionable '
    'for educators, students, and citizen scientists worldwide, while simultaneously generating valuable '
    'ground-truth data that improves the accuracy and reliability of satellite-based atmospheric monitoring.'
)

# --- Footer ---
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Report —')
run.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Sources: globe.gov | mynasadata.larc.nasa.gov | nesec.strategies.org')
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(128, 128, 128)

# Save
output_path = r'd:\_NASA_F9_Breath\Claude_project\NASA_Air_Quality_Education_Efforts_Detailed_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
