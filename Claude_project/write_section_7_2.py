"""Generate Section 7.2 for the BREATHE proposal as a Word document."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.25

# --- Header ---
header_para = doc.add_paragraph()
run = header_para.add_run('BREATHE: ')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run2 = header_para.add_run('BR')
run2.bold = True
run2.underline = True
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'
run3 = header_para.add_run('idging ')
run3.bold = True
run3.font.size = Pt(12)
run3.font.name = 'Times New Roman'
run4 = header_para.add_run('E')
run4.bold = True
run4.underline = True
run4.font.size = Pt(12)
run4.font.name = 'Times New Roman'
run5 = header_para.add_run('ducation on ')
run5.bold = True
run5.font.size = Pt(12)
run5.font.name = 'Times New Roman'
run6 = header_para.add_run('A')
run6.bold = True
run6.underline = True
run6.font.size = Pt(12)
run6.font.name = 'Times New Roman'
run7 = header_para.add_run('tmosphere, ')
run7.bold = True
run7.font.size = Pt(12)
run7.font.name = 'Times New Roman'
run8 = header_para.add_run('T')
run8.bold = True
run8.underline = True
run8.font.size = Pt(12)
run8.font.name = 'Times New Roman'
run9 = header_para.add_run('echnology, and ')
run9.bold = True
run9.font.size = Pt(12)
run9.font.name = 'Times New Roman'
run10 = header_para.add_run('HE')
run10.bold = True
run10.underline = True
run10.font.size = Pt(12)
run10.font.name = 'Times New Roman'
run11 = header_para.add_run('alth')
run11.bold = True
run11.font.size = Pt(12)
run11.font.name = 'Times New Roman'

# Horizontal line
from docx.oxml.ns import qn
from lxml import etree
border_para = doc.add_paragraph()
border_para.paragraph_format.space_after = Pt(6)
pPr = border_para._element.get_or_add_pPr()
pBdr = etree.SubElement(pPr, qn('w:pBdr'))
bottom = etree.SubElement(pBdr, qn('w:bottom'))
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), '000000')

# --- Section heading ---
heading = doc.add_paragraph()
heading.paragraph_format.space_before = Pt(12)
heading.paragraph_format.space_after = Pt(6)
run = heading.add_run('7.2 Data Exploration Platform')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# --- Body text paragraphs ---
def add_body(text, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.25
    # Handle bold markers **text**
    import re
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'
        else:
            r = p.add_run(part)
            r.font.size = Pt(12)
            r.font.name = 'Times New Roman'
    return p

# Paragraph 1: Context — existing NASA platforms and the gap
add_body(
    'Several NASA-supported platforms already provide valuable access to atmospheric '
    'data for educational use. The GLOBE Program offers structured citizen science protocols '
    'and the GLOBE Observer app for collecting ground-based aerosol, cloud, and weather '
    'observations that validate satellite measurements. My NASA Data, a NASA Langley Research '
    'Center project, curates satellite-derived datasets into mini-lessons, interactives, and '
    'lesson plans organized by Earth system phenomena, including a dedicated air quality topic '
    'with resources spanning PM2.5 pattern analysis, dust transport, and wildfire smoke modeling. '
    'The NASA Earth Science Education Collaborative (NESEC) aggregates these resources into a '
    'searchable portfolio that connects educators to materials by topic, grade level, and mission. '
    'These platforms offer broad geographic and topical coverage and serve as foundational '
    'infrastructure for NASA Earth science education. However, they are not designed to let '
    'participants explore how satellite observations relate to the specific atmospheric conditions '
    'and health outcomes in their own communities, nor do they integrate the newest generation '
    'of air quality missions\u2014TEMPO, PACE, and MAIA\u2014into a single, locally contextualized, '
    'health-connected learning environment.'
)

# Paragraph 2: What BREATHE platform does differently
add_body(
    'BREATHE addresses this gap through a purpose-built **Data Exploration Platform** that '
    'provides an online interface for curated NASA datasets, interactive visualization tools, '
    'and guided investigations anchored in participants\u2019 own communities. Unlike existing tools '
    'that present data at continental or global scales, the BREATHE platform is organized around '
    '**place-based entry points**: users begin by selecting or searching for their community '
    '(e.g., Houston Ship Channel, Rio Grande Valley, El Paso border corridor) and immediately '
    'see satellite observations, environmental context, and health indicators relevant to that '
    'location. This design reflects the project\u2019s core premise that meaningful engagement with '
    'NASA data requires connecting satellite observations to conditions participants experience '
    'in daily life.'
)

# Paragraph 3: Technical architecture and data layers
add_body(
    'The platform integrates multiple NASA data layers into a unified, map-based explorer. '
    'Users can toggle and overlay **TEMPO hourly NO\u2082, ozone, and formaldehyde** data to examine '
    'diurnal air pollution variability; **MAIA PM2.5 concentration estimates** to visualize '
    'particulate matter exposure at community scales; **MODIS/VIIRS aerosol optical depth, active '
    'fire detections, and land surface temperature** for long-term context and extreme event '
    'tracking; and **HAQAST satellite-driven PM2.5, organic carbon, and elemental carbon surface '
    'concentration estimates** at census-tract resolution. A compare mode allows side-by-side '
    'examination of two locations or time periods, and a temporal slider enables users to animate '
    'changes across hours (TEMPO), days (PACE and MAIA), or seasons (MODIS/VIIRS), reinforcing '
    'how atmospheric conditions vary across time and space. Each data layer includes embedded '
    'explanatory context\u2014\u201cWhat Am I Seeing?\u201d panels that describe the measurement technique, '
    'units, and health significance in accessible language\u2014so that the visualization serves as '
    'both a data tool and a learning scaffold.'
)

# Paragraph 4: Health data integration — key differentiator
add_body(
    'A distinguishing feature of the BREATHE platform is its integration of **community-level '
    'health indicators** alongside atmospheric data. During Stage 1, the platform connects NASA '
    'science assets to de-identified, ZIP-code-aggregated health records from regional healthcare '
    'information exchanges (the Paso del Norte HIE for El Paso and the Rio Grande Valley HIE for '
    'Brownsville) and Texas Inpatient, Outpatient, and Emergency Department Data for Harris County '
    '(Houston), with visualizations focused on cardiopulmonary outcomes such as asthma, acute '
    'bronchitis, upper respiratory infection emergency department visits, and COPD and '
    'cardiovascular disease incidence. For Stages 2 and 3, the platform scales this capability '
    'nationally by incorporating census-tract-level health outcome and social vulnerability '
    'indicators from the **CDC PLACES** dataset and the **CDC/ATSDR Social Vulnerability Index (SVI)**. '
    'This layered design enables participants to examine, for example, how a neighborhood with '
    'elevated TEMPO NO\u2082 concentrations also shows higher asthma-related emergency department '
    'visits and greater social vulnerability\u2014making the connection between atmospheric science '
    'and public health tangible and data-driven. No existing NASA education platform provides '
    'this satellite-to-health integration.'
)

# Paragraph 5: Scaffolded engagement and modules
add_body(
    'The platform organizes learning experiences along a four-level **engagement progression** '
    'aligned with the NASA Public-Science Engagement Spectrum: (1) **Awareness** modules (e.g., '
    '\u201cWhat\u2019s in the Air?\u201d) introduce satellite-based air quality concepts through guided '
    'exploration of NASA Worldview and TEMPO imagery; (2) **Learning** modules (e.g., \u201cPollution '
    'Around the Clock,\u201d \u201cSmoke Signals from Space\u201d) build data interpretation skills through '
    'structured investigations of hourly TEMPO patterns and wildfire smoke transport; '
    '(3) **Experience** modules (e.g., \u201cParticles That Matter,\u201d \u201cUrban Heat and Health\u201d) engage '
    'participants in hands-on analysis of MAIA PM2.5 data and environmental justice dimensions; '
    'and (4) **Contribution** modules (e.g., \u201cBe a Community Air Monitor\u201d) guide participants '
    'through building low-cost PM2.5 sensors, collecting ground measurements, and comparing '
    'them with satellite observations to validate NASA data. Each module specifies grade level '
    '(7\u201312), duration, NGSS alignment, learning objectives, and the NASA data assets used, and '
    'all materials are available in English and Spanish.'
)

# Paragraph 6: Regional data stories and educator hub
add_body(
    'The platform also features **regional data stories** that narrate environmental challenges '
    'specific to each pilot community\u2014petrochemical emissions and hurricane-related chemical '
    'releases in Houston, transboundary pollution and agricultural dust along the U.S.\u2013Mexico '
    'border, cross-border emissions and desert dust in El Paso\u2014using curated NASA data '
    'visualizations and locally relevant health context. An integrated **Educator Hub** provides '
    'classroom-ready lesson plans, data activities aligned with the Texas Essential Knowledge and '
    'Skills (TEKS) and NGSS standards, and direct links to professional development resources '
    '(Section 7.3). Educators can filter modules by grade level, topic, engagement level, and '
    'NASA mission, and access downloadable investigation worksheets and discussion guides. By '
    'building on the strengths of GLOBE, My NASA Data, and NESEC\u2014and linking directly to their '
    'resources where appropriate\u2014while adding place-based health integration, scaffolded '
    'engagement, and next-generation mission data that these platforms do not yet offer, the '
    'BREATHE Data Exploration Platform fills a critical gap in NASA\u2019s Science Activation '
    'ecosystem.'
)

# Save
output_path = r'd:\_NASA_F9_Breath\Claude_project\BREATHE_Section_7_2_Data_Exploration_Platform.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
