"""Generate Sections 7.1 and 7.2 for the BREATHE proposal as a Word document."""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from lxml import etree
import re

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.25


def add_breathe_header():
    header_para = doc.add_paragraph()
    parts = [
        ('BREATHE: ', True, False),
        ('BR', True, True),
        ('idging ', True, False),
        ('E', True, True),
        ('ducation on ', True, False),
        ('A', True, True),
        ('tmosphere, ', True, False),
        ('T', True, True),
        ('echnology, and ', True, False),
        ('HE', True, True),
        ('alth', True, False),
    ]
    for text, bold, underline in parts:
        r = header_para.add_run(text)
        r.bold = bold
        r.underline = underline
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'

    # Horizontal line
    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_after = Pt(6)
    pPr = border_para._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')


def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    return p


def add_body(text, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.25
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'
    return p


# ============================================================
# SECTION 7.1 — NASA Earth Science Learning Modules
# ============================================================
add_breathe_header()
add_section_heading('7.1 NASA Earth Science Learning Modules')

add_body(
    'BREATHE develops a suite of learning modules that use curated datasets from TEMPO, PACE, '
    'MAIA, and HAQAST data fusion products (Section 4) to explore atmospheric processes, air '
    'pollution patterns, and environmental exposure in locally relevant contexts. Each module '
    'is a self-contained, classroom-ready investigation that positions participants as active '
    'analysts of NASA EO data rather than passive recipients of content. Modules target grades '
    '7\u201312, include educator facilitation guides, and are available in English and Spanish. '
    'The module suite is organized along a four-level **engagement progression** aligned with '
    'the NASA Public-Science Engagement Spectrum [NASA Science Mission Directorate, 2025]:'
)

add_body(
    '**Awareness.** Introductory modules such as \u201cWhat\u2019s in the Air?\u201d (45\u201360 minutes, '
    'grades 6\u20138) introduce key air pollutants and their health effects, explain how satellites '
    'measure atmospheric composition, and guide participants through NASA Worldview and TEMPO '
    'imagery to compare air quality across communities. '
    '**Learning.** \u201cPollution Around the Clock\u201d (60\u201390 minutes, grades 8\u201312) uses TEMPO\u2019s '
    'hourly NO\u2082 to investigate diurnal air quality variation, asking participants to identify '
    'emission sources, analyze meteorological modulation, and propose evidence-based '
    'recommendations for reducing exposure. \u201cSmoke Signals from Space\u201d (60\u201390 minutes, '
    'grades 7\u201312) uses MODIS/VIIRS fire detections and FIRMS data to trace a real wildfire '
    'smoke event from ignition to community health impact. '
    '**Experience.** \u201cParticles That Matter\u201d (two class periods, grades 9\u201312) guides '
    'participants through interpreting MAIA PM2.5 maps and comparing satellite-derived estimates '
    'with EPA AirNow ground measurements, extending to environmental justice dimensions by '
    'overlaying SVI data to analyze PM exposure across communities with different socioeconomic '
    'characteristics. \u201cUrban Heat and Health\u201d (90 minutes, grades 9\u201312) maps heat exposure '
    'using MODIS/VIIRS land surface temperature and Landsat thermal data. '
    '**Contribution.** \u201cBe a Community Air Monitor\u201d (4\u20136 hours, grades 8\u201312 and community '
    'participants) guides learners through the full citizen science cycle: assembling low-cost '
    'PM2.5 sensors, collecting and quality-checking measurements, comparing ground-based data '
    'with satellite observations, and communicating findings to stakeholders. This module links '
    'to the citizen science activities in Section 7.4 and produces data integrated into the '
    'Data Exploration Platform (Section 7.2).'
)

add_body(
    'All modules specify NGSS and TEKS alignment, learning objectives, NASA data assets, and '
    'estimated duration. They complement\u2014rather than duplicate\u2014existing SciAct resources: '
    'where GLOBE provides protocols for atmospheric measurement and My NASA Data offers '
    'mini-lessons using historical datasets, BREATHE modules integrate next-generation mission '
    'data these platforms do not yet feature, anchor investigations in local environmental '
    'conditions, and explicitly connect atmospheric observations to community health outcomes. '
    'Modules are iteratively refined through co-design with educators (Section 7.3) and '
    'feedback from pilot implementation.'
)

# ============================================================
# PAGE BREAK and SECTION 7.2
# ============================================================
doc.add_page_break()
add_breathe_header()
add_section_heading('7.2 Data Exploration Platform')

add_body(
    'Existing SciAct platforms\u2014GLOBE, My NASA Data, NESEC\u2014provide valuable but '
    'general-purpose access to atmospheric data. They are not designed to let participants '
    'explore how satellite observations relate to specific atmospheric conditions and health '
    'outcomes in their own communities, nor do they integrate TEMPO, PACE, and MAIA alongside '
    'community-collected citizen science data into a single, locally contextualized environment. '
    'BREATHE addresses this gap through a purpose-built **Data Exploration Platform** that '
    'unifies curated NASA satellite datasets, citizen science observations from BREATHE '
    'participants (Section 7.4), and community-level health indicators within an interactive '
    'visualization interface. Unlike existing tools that present data at continental scales, '
    'the platform is organized around **place-based entry points**: users select their community '
    '(e.g., Houston Ship Channel, Rio Grande Valley, El Paso) and immediately see satellite '
    'observations, ground measurements, and health indicators relevant to that location.'
)

add_body(
    'The platform provides a **map-based explorer** with toggleable data layers drawn from the '
    'NASA science assets described in Section 4: TEMPO hourly NO\u2082, ozone, and formaldehyde; '
    'PACE aerosol optical depth and particle composition; MAIA PM2.5 mass and speciated '
    'concentrations; MODIS/VIIRS aerosol optical depth, fire detections, and land surface '
    'temperature; and HAQAST surface PM2.5, organic carbon, and elemental carbon estimates at '
    'census-tract resolution. Alongside these satellite products, the platform displays '
    '**citizen science data**\u2014ground-level PM2.5 from low-cost sensors at schools and community '
    'sites, sky condition observations, and locally recorded environmental conditions (Section '
    '7.4). Co-locating satellite and ground-based data on the same interface enables direct '
    'comparison between what NASA observes from orbit and what participants measure on the '
    'ground, reinforcing data validation concepts and making citizen science contributions '
    'scientifically meaningful. Interactive features include a **compare mode** for side-by-side '
    'examination of locations or time periods, a **temporal slider** animating changes across '
    'hours (TEMPO), days (PACE, MAIA), or seasons (MODIS/VIIRS), and **quick-view presets** for '
    'pilot regions. Each layer includes explanatory panels describing measurement technique, '
    'units, and health significance, so the visualization serves as both a data tool and a '
    'learning scaffold.'
)

add_body(
    'A distinguishing feature is the integration of **community-level health indicators** '
    'alongside atmospheric and citizen science data. As described in Section 4, during Stage 1 '
    'the platform connects satellite and ground-based observations to de-identified health '
    'records from regional HIEs and Texas hospital data, with visualizations focused on '
    'cardiopulmonary outcomes such as asthma and respiratory infection ED visits. For Stages 2 '
    'and 3, the platform scales nationally using CDC PLACES and SVI indicators. This layered '
    'design enables learners to observe, for instance, that a neighborhood with elevated TEMPO '
    'NO\u2082 and high PM2.5 from a BREATHE community sensor also shows above-average asthma ED '
    'visits and greater social vulnerability\u2014making the satellite-to-ground-to-health connection '
    'tangible and data-driven. The platform also features **regional data stories** narrating '
    'environmental challenges specific to each pilot community, and an integrated **Educator '
    'Hub** with classroom-ready activities aligned with TEKS and NGSS, filterable by grade level, '
    'topic, and NASA mission. By building on GLOBE, My NASA Data, and NESEC while adding '
    'place-based health integration, citizen science co-visualization, and next-generation '
    'mission data, the BREATHE platform fills a critical gap in NASA\u2019s SciAct ecosystem.'
)

# Save
output_path = r'd:\_NASA_F9_Breath\Claude_project\BREATHE_Sections_7_1_and_7_2.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
